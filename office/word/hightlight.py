from docx import Document
from docx.shared import RGBColor

# 加载 Word 文档
doc = Document(r'D:\workspace\myProjects\word_cmp\static\Black Beauty 原文.docx')

# 要高亮的单词
highlight_word = "what"

# 遍历文档中的段落
for para in doc.paragraphs:
    if highlight_word in para.text:
        # 分割文本，逐个单词处理
        words = para.text.split()
        new_run = para.clear()  # 清空原段落
        for word in words:
            run = para.add_run(word + ' ')  # 添加原单词和空格
            if word == highlight_word:
                # 设置背景颜色为黄色（模拟高亮）
                run.font.color.rgb = RGBColor(255, 200, 100)  # 例如黄色

# 保存修改后的文档
doc.save('highlighted_document.docx')
