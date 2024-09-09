import docx
from docx.shared import RGBColor, Pt


def set_run(run, font_size, bold, color):
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color


file_name = "2.docx"  # 替换为你的文档文件名
doc = docx.Document(file_name)

for para in doc.paragraphs:
    if "NS1和" in para.text:
        text_parts = para.text.split("NS1和")
        para.clear()  # 清空原来的文本
        for i, text_part in enumerate(text_parts):
            run = para.add_run(text_part)
            font_size = run.font.size
            bold = run.bold
            color = run.font.color.rgb
            set_run(run, font_size, bold, color)
            if i < len(text_parts) - 1:
                run = para.add_run("NS1和")
                set_run(run, font_size, bold, RGBColor(255, 0, 0))  # 将红色设置为RGBColor(255, 0, 0)
doc.save("23.docx")  # 保存处理后的文档

